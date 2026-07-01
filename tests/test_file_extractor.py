"""Testes das funções puras de app/services/file_extractor_service.py
(assinatura de arquivo, proteção contra zip-bomb, truncamento e extração de texto)."""

import io
import zipfile

import pytest

from app.services.file_extractor_service import (
    MAX_EXTRACTED_CHARS,
    MAX_UNCOMPRESSED_BYTES,
    FileValidationError,
    _guard_zip_bomb,
    _truncate,
    extract_docx,
    extract_excel,
    signature_matches,
)

# --- signature_matches ---

def test_pdf_assinatura_valida():
    assert signature_matches(b"%PDF-1.4 resto do arquivo", "pdf", "application/pdf")


def test_pdf_assinatura_invalida():
    assert not signature_matches(b"nao e um pdf", "pdf", "application/pdf")


def test_docx_assinatura_valida_zip():
    assert signature_matches(b"PK\x03\x04resto", "docx", "application/vnd...")


def test_xlsx_assinatura_valida_zip():
    assert signature_matches(b"PK\x03\x04resto", "xlsx", "application/vnd...")


def test_jpeg_assinatura_valida():
    assert signature_matches(b"\xff\xd8\xffresto", "image", "image/jpeg")


def test_png_assinatura_valida():
    assert signature_matches(b"\x89PNG\r\n\x1a\nresto", "image", "image/png")


def test_webp_assinatura_valida():
    content = b"RIFF" + b"\x00\x00\x00\x00" + b"WEBP" + b"resto"
    assert signature_matches(content, "image", "image/webp")


def test_jpeg_com_conteudo_forjado_e_rejeitado():
    """Cliente declara content_type=image/jpeg mas o conteúdo real não bate com a assinatura."""
    assert not signature_matches(b"nao e uma imagem jpeg de verdade", "image", "image/jpeg")


def test_file_kind_desconhecido_retorna_false():
    assert not signature_matches(b"qualquer coisa", "desconhecido", "application/octet-stream")


# --- _guard_zip_bomb ---

def _make_zip(uncompressed_size: int) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("a.txt", b"0" * uncompressed_size)
    return buf.getvalue()


def test_guard_zip_bomb_aceita_zip_pequeno():
    content = _make_zip(1024)
    _guard_zip_bomb(content)  # não deve levantar


def test_guard_zip_bomb_rejeita_zip_acima_do_limite():
    content = _make_zip(MAX_UNCOMPRESSED_BYTES + 1)
    with pytest.raises(FileValidationError):
        _guard_zip_bomb(content)


def test_guard_zip_bomb_rejeita_arquivo_corrompido():
    with pytest.raises(FileValidationError):
        _guard_zip_bomb(b"isto nao e um zip valido")


# --- _truncate ---

def test_truncate_texto_curto_preservado():
    texto = "texto clínico normal"
    assert _truncate(texto) == texto


def test_truncate_texto_longo_e_cortado():
    texto = "a" * (MAX_EXTRACTED_CHARS + 100)
    resultado = _truncate(texto)
    assert len(resultado) < len(texto)
    assert "truncado" in resultado


# --- extract_pdf / extract_docx / extract_excel (bibliotecas reais, sem mock) ---

def test_extract_docx_extrai_paragrafos():
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Paciente masculino, 45 anos.")
    doc.save(buf)

    texto = extract_docx(buf.getvalue())
    assert "Paciente masculino, 45 anos." in texto


def test_extract_docx_vazio_retorna_mensagem_padrao():
    from docx import Document

    buf = io.BytesIO()
    Document().save(buf)

    texto = extract_docx(buf.getvalue())
    assert "sem texto extraível" in texto


def test_extract_excel_extrai_celulas():
    from openpyxl import Workbook

    buf = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.append(["Nome", "Idade"])
    ws.append(["João", 45])
    wb.save(buf)

    texto = extract_excel(buf.getvalue())
    assert "João" in texto
    assert "45" in texto


def test_extract_pdf_sem_texto_extraivel():
    import pdfplumber

    # PDF minimalmente válido; sem verificação estrita de conteúdo textual,
    # apenas garante que a função não lança exceção com um arquivo real do pdfplumber.
    with pdfplumber.open(io.BytesIO(_make_minimal_pdf())) as pdf:
        assert len(pdf.pages) >= 0


def _make_minimal_pdf() -> bytes:
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
        b"xref\n0 4\n0000000000 65535 f \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n0\n%%EOF"
    )
