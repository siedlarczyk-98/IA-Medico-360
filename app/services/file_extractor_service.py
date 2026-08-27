"""
Médico 360 — Extração de texto de arquivos enviados pelo médico.

Suporta: PDF, DOCX, XLSX, e imagens (JPEG/PNG/WEBP via Claude Haiku vision).

Os parsers de documento (PDF/DOCX/XLSX) são síncronos e CPU-bound — devem ser
chamados a partir de uma thread (anyio.to_thread.run_sync) para não bloquear o
event loop. Há tetos durante a extração e proteção contra zip-bombs (DOCX/XLSX
são arquivos ZIP que poderiam expandir para gigabytes).
"""

import base64
import io
import zipfile
from uuid import UUID

from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.http_client import get_client
from app.middleware.dlp import sanitize_prompt_async
from app.models.models import FileExtraction

settings = get_settings()

ALLOWED_CONTENT_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "image/jpeg": "image",
    "image/png": "image",
    "image/webp": "image",
}

MAX_FILE_BYTES = 10 * 1024 * 1024   # 10 MB (geral)
MAX_IMAGE_BYTES = 5 * 1024 * 1024   # 5 MB (limite base64 Anthropic)
MAX_EXTRACTED_CHARS = 50_000

# ── Limites anti-DoS durante a extração ──────────────────────
MAX_PDF_PAGES = 100
MAX_DOCX_PARAGRAPHS = 20_000
MAX_DOCX_TABLE_ROWS = 20_000
MAX_XLSX_ROWS = 50_000
# Teto de tamanho descomprimido para arquivos ZIP (DOCX/XLSX) — barra zip-bombs.
MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024  # 200 MB


class FileValidationError(Exception):
    """Erro de validação de conteúdo do arquivo (assinatura, zip-bomb, etc.)."""


# Marca gravada quando o PDF não tem camada de texto — tipicamente um laudo
# digitalizado. Constante, e não literal solto, porque o endpoint precisa
# reconhecer o caso para avisar o médico: antes ele anexava, recebia uma
# resposta pobre e não tinha como saber por quê.
PDF_SEM_TEXTO = "(PDF sem texto extraível — pode ser escaneado sem OCR)"

AVISO_PDF_ESCANEADO = (
    "Este PDF parece ser digitalizado: não foi possível ler texto dele. "
    "Para que o exame seja analisado, envie a página como imagem (JPG ou PNG) — "
    "esse caminho usa leitura visual."
)

AVISO_ARQUIVO_VAZIO = (
    "Não foi possível extrair texto deste arquivo. Confira se ele não está "
    "vazio ou protegido."
)


def aviso_de_extracao(texto: str, file_type: str) -> str | None:
    """
    Aviso a mostrar ao médico quando a extração não rendeu conteúdo útil.

    Devolve None no caminho normal. Existe porque falhar em silêncio aqui é
    especialmente ruim: o arquivo é aceito, a mensagem é enviada, e o modelo
    responde sobre um exame que nunca chegou até ele.
    """
    if file_type == "image":
        return None  # imagem não passa por extração de texto

    if texto == PDF_SEM_TEXTO:
        return AVISO_PDF_ESCANEADO
    if not texto or not texto.strip():
        return AVISO_ARQUIVO_VAZIO
    return None


# ── Resolução de contexto de arquivo ─────────────────────────

async def resolve_file_context(
    prompt: str,
    file_id: UUID | None,
    user_id: UUID,
    db: AsyncSession,
    *,
    support_vision: bool = True,
) -> tuple[str, dict | None]:
    """
    Resolve o contexto de uma extração de arquivo previamente enviada.

    Retorna (prompt enriquecido, image_content | None). Checa a posse do arquivo
    (user_id) — levanta 404 se não pertencer ao usuário.

    Quando `support_vision=True` e o arquivo é imagem, devolve `image_content` com
    os pixels (base64) para envio como vision block; a descrição já sanitizada vai
    como `fallback_text` para providers sem visão. Quando `support_vision=False`,
    imagens são tratadas como texto (apenas a descrição é injetada no prompt).
    """
    if not file_id:
        return prompt, None

    extraction = await db.get(FileExtraction, file_id)
    if not extraction or extraction.user_id != user_id:
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")

    if support_vision and extraction.file_type == "image" and extraction.image_base64:
        enriched = f"[Imagem: {extraction.file_name}]\n\n{prompt}".strip()
        image_content = {
            "base64": extraction.image_base64,
            "media_type": extraction.image_media_type,
            "fallback_text": extraction.extracted_text,
        }
        return enriched, image_content

    enriched = f"[Arquivo: {extraction.file_name}]\n{extraction.extracted_text}\n\n---\n\n{prompt}".strip()
    return enriched, None


# Teto de anexos por mensagem. Cada imagem custa uma chamada de visão na
# extração e pesa base64 dentro do prompt; sem teto, uma mensagem com vinte
# arquivos estoura o contexto e a conta.
MAX_ANEXOS_POR_MENSAGEM = 5


async def resolve_files_context(
    prompt: str,
    file_ids: list[UUID],
    user_id: UUID,
    db: AsyncSession,
    *,
    support_vision: bool = True,
) -> tuple[str, list[dict], list[FileExtraction]]:
    """
    Resolve VÁRIOS anexos de uma mensagem.

    Devolve (prompt enriquecido, imagens para envio como visão, extrações).

    Cada arquivo tem a posse checada individualmente — misturar um `file_id`
    alheio no meio de arquivos próprios não passa.

    A ordem importa: os textos entram antes da pergunta, na ordem em que o
    médico anexou, para que "compare o primeiro com o segundo" tenha sentido.
    """
    if not file_ids:
        return prompt, [], []

    if len(file_ids) > MAX_ANEXOS_POR_MENSAGEM:
        raise HTTPException(
            status_code=422,
            detail=f"Máximo de {MAX_ANEXOS_POR_MENSAGEM} arquivos por mensagem.",
        )

    extractions: list[FileExtraction] = []
    for file_id in file_ids:
        extraction = await db.get(FileExtraction, file_id)
        if not extraction or extraction.user_id != user_id:
            raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
        extractions.append(extraction)

    imagens: list[dict] = []
    blocos_texto: list[str] = []

    for extraction in extractions:
        e_imagem = extraction.file_type == "image" and extraction.image_base64
        if support_vision and e_imagem:
            imagens.append({
                "base64": extraction.image_base64,
                "media_type": extraction.image_media_type,
                "fallback_text": extraction.extracted_text,
                "file_name": extraction.file_name,
            })
            # O nome entra no texto mesmo quando a imagem vai como visão: sem
            # isso o modelo não consegue dizer "na primeira tomografia..." de
            # forma que o médico saiba a qual arquivo ele se refere.
            blocos_texto.append(f"[Imagem anexada: {extraction.file_name}]")
        elif e_imagem:
            # Provider sem visão: sobra a descrição gerada na extração, e o
            # modelo precisa saber que está lendo descrição, não a imagem.
            blocos_texto.append(
                f"[Descrição da imagem {extraction.file_name} — a imagem em si não foi enviada a este modelo]\n"
                f"{extraction.extracted_text}"
            )
        else:
            blocos_texto.append(f"[Arquivo: {extraction.file_name}]\n{extraction.extracted_text}")

    enriched = "\n\n".join([*blocos_texto, "---", prompt]).strip()
    return enriched, imagens, extractions


# ── Validação de assinatura (magic bytes) ────────────────────

def signature_matches(content: bytes, file_kind: str, content_type: str) -> bool:
    """Confere a assinatura real do arquivo contra o tipo declarado pelo cliente."""
    head = content[:16]
    if file_kind == "pdf":
        return head.startswith(b"%PDF")
    if file_kind in ("docx", "xlsx"):
        # OOXML são arquivos ZIP. PK\x03\x04 (normal), PK\x05\x06 (vazio), PK\x07\x08 (spanned).
        return head[:2] == b"PK"
    if file_kind == "image":
        if content_type == "image/jpeg":
            return head.startswith(b"\xff\xd8\xff")
        if content_type == "image/png":
            return head.startswith(b"\x89PNG\r\n\x1a\n")
        if content_type == "image/webp":
            return head[:4] == b"RIFF" and head[8:12] == b"WEBP"
    return False


def _guard_zip_bomb(content: bytes) -> None:
    """Aborta se a soma dos tamanhos descomprimidos do ZIP exceder o teto."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            total = sum(info.file_size for info in zf.infolist())
    except zipfile.BadZipFile:
        raise FileValidationError("Arquivo corrompido ou não é um documento válido.")
    if total > MAX_UNCOMPRESSED_BYTES:
        raise FileValidationError("Arquivo excede o limite de descompressão permitido.")


def _truncate(text: str) -> str:
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS] + "\n\n[... conteúdo truncado — arquivo muito extenso]"
    return text


def extract_pdf(content: bytes) -> str:
    import pdfplumber  # import tardio — não penaliza boot se lib ausente

    parts: list[str] = []
    acc_len = 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages[:MAX_PDF_PAGES]:
            text = page.extract_text()
            if text and text.strip():
                stripped = text.strip()
                parts.append(stripped)
                acc_len += len(stripped)
                if acc_len > MAX_EXTRACTED_CHARS:
                    break
    joined = "\n\n".join(parts) if parts else PDF_SEM_TEXTO
    return _truncate(joined)


def extract_docx(content: bytes) -> str:
    from docx import Document  # python-docx

    _guard_zip_bomb(content)
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs[:MAX_DOCX_PARAGRAPHS]:
        if para.text.strip():
            parts.append(para.text.strip())

    row_count = 0
    for table in doc.tables:
        for row in table.rows:
            if row_count >= MAX_DOCX_TABLE_ROWS:
                break
            row_count += 1
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
        if row_count >= MAX_DOCX_TABLE_ROWS:
            break

    joined = "\n".join(parts) if parts else "(Documento sem texto extraível)"
    return _truncate(joined)


def extract_excel(content: bytes) -> str:
    import openpyxl

    _guard_zip_bomb(content)
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []
    row_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            if row_count >= MAX_XLSX_ROWS:
                break
            row_count += 1
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Planilha: {sheet_name}]\n" + "\n".join(rows))
        if row_count >= MAX_XLSX_ROWS:
            break

    wb.close()
    joined = "\n\n".join(parts) if parts else "(Planilha sem dados extraíveis)"
    return _truncate(joined)


async def extract_image(content: bytes, media_type: str) -> dict:
    """
    Retorna dict com:
      - base64: str  → bytes da imagem em base64 (para envio direto como vision block)
      - media_type: str → MIME type original (ex: "image/jpeg")
      - description: str → descrição gerada pelo Claude Haiku (fallback para modelos sem visão).
        A descrição passa pelo DLP antes de ser persistida/retornada — ela pode conter
        PII transcrita da imagem e é injetada no prompt de providers sem visão (Perplexity).
    """
    b64 = base64.standard_b64encode(content).decode()
    payload = {
        "model": "claude-haiku-4-5-20251001",
        "max_tokens": 1024,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": media_type, "data": b64},
                    },
                    {
                        "type": "text",
                        "text": (
                            "Você é um assistente médico. Analise esta imagem e extraia todo o conteúdo "
                            "relevante em texto estruturado.\n"
                            "- Se for exame laboratorial: liste todos os itens, valores e referências.\n"
                            "- Se for imagem diagnóstica (Rx, TC, RM, US): descreva os achados visíveis.\n"
                            "- Se for prescrição ou receita: liste medicamentos, doses e posologia.\n"
                            "- Se for outro documento médico: transcreva o conteúdo.\n"
                            "Seja preciso e objetivo. Não adicione interpretações clínicas."
                        ),
                    },
                ],
            }
        ],
    }

    client = get_client()
    resp = await client.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": settings.anthropic_api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json=payload,
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    description = data["content"][0]["text"]
    # DLP: a descrição vira fallback_text injetado em providers sem visão (ex.: Perplexity),
    # fora do fluxo de sanitização do prompt — então sanitizamos aqui.
    description = (await sanitize_prompt_async(description)).sanitized_text
    usage = data.get("usage", {})
    return {
        "base64": b64,
        "media_type": media_type,
        "description": description,
        "tokens_in": usage.get("input_tokens", 0),
        "tokens_out": usage.get("output_tokens", 0),
    }
